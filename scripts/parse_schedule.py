# -*- coding: utf-8 -*-
"""
课程表解析脚本
==============
从郑州大学学生课表 docx 中提取"课程 + 老师"映射，去重后输出 JSON。

用法：
    python scripts/parse_schedule.py <课表.docx路径> [输出.json路径]

输出格式：
{
  "student": {"name": "吴冠辉", "grade": "2023", "major": "物联网工程", "class": "物联网工程2023级3班"},
  "courses": [
    {"course": "RFID原理及应用", "teacher": "张博"},
    ...
  ],
  "teachers": ["张博", "王有为", ...]   # 去重后的老师名单
}
"""

import re
import sys
import json
import os

try:
    from docx import Document
except ImportError:
    print("缺少 python-docx，请先安装：pip install python-docx")
    sys.exit(1)


# 教室号模式：北区9_102 / 北区6_304（图形与视觉计算）
ROOM_PATTERN = r"(?:北区)?[\d_]+(?:（[^）]+）)?"
# 老师名：教室号后紧跟 2-3 个中文字（用 \xa0 或空格分隔）
TEACHER_PATTERN = re.compile(
    r"北区[\d_]+（[^）]+）\s*[\xa0 ]+([\u4e00-\u9fff]{2,3})|北区[\d_]+[\u4e00-\u9fff]?\s*[\xa0 ]+([\u4e00-\u9fff]{2,3})"
)


def extract_student_info(doc):
    """从段落中提取学生信息（姓名、年级、专业、班级）"""
    info = {"name": "", "grade": "", "major": "", "class": ""}
    for p in doc.paragraphs:
        text = p.text.strip()
        if "年级" in text and "姓名" in text:
            grade_m = re.search(r"年级[:：]\s*(\d+)", text)
            major_m = re.search(r"专业[:：]\s*([^\s]+)", text)
            class_m = re.search(r"班级[:：]\s*([^\s]+)", text)
            name_m = re.search(r"姓名[:：]\s*([^\s]+)", text)
            if grade_m: info["grade"] = grade_m.group(1)
            if major_m: info["major"] = major_m.group(1)
            if class_m: info["class"] = class_m.group(1)
            if name_m: info["name"] = name_m.group(1)
            break
    return info


def extract_courses(doc):
    """遍历所有表格单元格，提取课程-老师映射"""
    pairs = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if not text.strip():
                    continue
                # 整个单元格作为文本：按课程块特征切分
                # 每个课程块以"课程名\n课程号\n..."开头，老师名在教室号后
                # 用正则找所有 (课程名, 老师名) 组合
                for m in TEACHER_PATTERN.finditer(text):
                    teacher = (m.group(1) or m.group(2)).strip()
                    # 课程名：在老师匹配位置之前，找块内第一行的第一个词
                    head = text[:m.start()]
                    # 取最后一个 "课程名\n课程号" 块
                    blocks = re.split(r"[\n]", head)
                    course = ""
                    for b in reversed(blocks):
                        b = b.strip()
                        # 跳过纯数字块（课程号 773308、391015.Y2、人数 97 等）
                        if re.match(r"^[\d.]+", b):
                            continue
                        cm = re.match(r"^[\u4e00-\u9fffA-Za-z0-9]+", b)
                        if cm:
                            course = cm.group(0)
                            break
                    if teacher and course:
                        pairs.append({"course": course, "teacher": teacher})
    return pairs


def dedupe(pairs):
    """去重课程-老师对"""
    seen = set()
    result = []
    for pair in pairs:
        key = (pair["course"], pair["teacher"])
        if key not in seen:
            seen.add(key)
            result.append(pair)
    return result


def parse_schedule_file(docx_path: str) -> dict:
    """
    解析课表 docx 文件，返回结构化数据。
    供 API 调用（与命令行 main 共用同一套逻辑）。
    """
    doc = Document(docx_path)
    student = extract_student_info(doc)
    pairs = dedupe(extract_courses(doc))
    teachers = sorted(set(p["teacher"] for p in pairs))
    return {
        "student": student,
        "courses": pairs,
        "teachers": teachers,
    }


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    docx_path = sys.argv[1]
    out_path = sys.argv[2] if len(sys.argv) > 2 else "teachers.json"

    if not os.path.exists(docx_path):
        print("文件不存在:", docx_path)
        sys.exit(1)

    doc = Document(docx_path)
    student = extract_student_info(doc)
    pairs = dedupe(extract_courses(doc))
    teachers = sorted(set(p["teacher"] for p in pairs))

    data = {
        "student": student,
        "courses": pairs,
        "teachers": teachers,
    }

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"学生: {student.get('name', '未知')} ({student.get('class', '')})")
    print(f"课程-老师对: {len(pairs)} 条")
    print(f"老师名单: {teachers}")


if __name__ == "__main__":
    main()
